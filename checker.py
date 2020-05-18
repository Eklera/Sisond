import random
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
import docx
from zipfile import ZipFile
from xml.dom.minidom import parseString
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lxml import etree


def file_check_interface(file, mock=True):
    if mock:
        return mock_return()
    return check_file(file)


def mock_return():
    is_wrong = random.random() < 0.4
    res = [1 for x in range(5)]
    if is_wrong:
        for i in range(len(res)):
            if random.random() < 0.5:
                res[i] = 2
    return res


# ======================================================================
def check_file(file_path):
    ans = [0, 0, 0, 0, 0]
    file = docx.Document(file_path)

    return ans


def xml_compare_theme_latin_font(xml_theme, font):
    if xml_theme is None:
        raise AttributeError("Нет xml файла")

    for x in xml_theme:
        for i in x:
            if "fontScheme" in i.tag:
                xml_theme = i
                break

    if xml_theme is None:
        raise AttributeError("Нет fontScheme в xml файле")

    theme_fonts = []
    for elem in xml_theme:
        for font_entry in elem:
            if "latin" in font_entry.tag:
                theme_fonts.append(font_entry)

    for theme_font in theme_fonts:
        if theme_font.attrib["typeface"] != font:
            return False
    return True


def get_default_font_is_ok(file):
    """
    :param file: .docx file to analyse
    :return: True if document body does not contain any text formatted with default theme or default theme font is Time. Else returns False.
    """
    # TRUE IF OK
    theme_font_is_times = False
    theme = None
    with ZipFile(file, 'r') as z:
        with z.open('word/theme/theme1.xml') as theme:  # todo: а я говорил
            theme = etree.parse(theme).getroot()
        document = z.read('word/document.xml')
        document = str(document)
    z.close()

    theme_font_is_times = xml_compare_theme_latin_font(theme, "Times New Roman")
    #
    # # todo: тоже парсить надо, они могут меняться местами
    # bad_string = ["asciiTheme=\"minorHAnsi\" w:hAnsiTheme=\"minorHAnsi\" w:cstheme=\"minorHAnsi\"",
    #               "asciiTheme=\"majorHAnsi\" w:hAnsiTheme=\"majorHAnsi\" w:cstheme=\"majorHAnsi\""]
    # doc_contains_default_theme = bad_string[0] in document or bad_string[1] in document
    #
    # return not doc_contains_default_theme \
    #        or doc_contains_default_theme and theme_font_is_times
    return theme_font_is_times


def get_theme_run(file):
    """
    :param file: name file to analyse
    :return: True if any run contain theme, else returns False
    """

    run_theme = {}
    with ZipFile(file, 'r') as z:
        with z.open('word/document.xml') as document:  # todo: а я говорил
            document = etree.parse(document).getroot()
    z.close()

    auto_content = False

    for x in document:
        for elem in x:
            if elem.tag.endswith('}p'):
                for p in elem:
                    if p.tag.endswith('pPr') or p.tag.endswith('}r'):
                        for pPr in p:
                            if pPr.tag.endswith('tabs'):
                                for tab in pPr:
                                    tab_attrib = tab.attrib
                                    for key in tab_attrib.keys():
                                        if key.endswith("leader"):
                                            auto_content = True
                            else:
                                auto_content = False
                            if pPr.tag.endswith('rPr') and auto_content is False:
                                for rPr in pPr:
                                    if rPr.tag.endswith('rFonts'):
                                        rPr_attrib = rPr.attrib
                                        for key in rPr_attrib.keys():
                                            # if key.endswith('cstheme'):
                                            #     run_theme["cstheme"] = rPr_attrib[key]
                                            #     print(rPr)
                                            if key.endswith('hAnsiTheme'):
                                                run_theme["hAnsiTheme"] = rPr_attrib[key]
                                            # elif key.endswith('eastAsiaTheme'):
                                            #     run_theme["eastAsiaTheme"] = rPr_attrib[key]
                                            elif key.endswith('asciiTheme'):
                                                run_theme["asciiTheme"] = rPr_attrib[key]

    if len(run_theme) != 0:
        return True
    else:
        return False


def get_styles_with_theme(file):
    """
    :param file: name file to analyse
    :return: dict with key: style id, value: theme
    """
    style_xml = None
    # style_font = {'cstheme': '',
    #               'hAnsiTheme': '',
    #               'eastAsiaTheme': '',
    #               'asciiTheme': ''}
    style_font = {}
    styles_with_theme = {}
    style_id = 0
    # styles_id = {style_id: style_font}
    with ZipFile(file, 'r') as z:
        with z.open('word/styles.xml') as style_xml:  # todo: а я говорил
            style_xml = etree.parse(style_xml).getroot()
    z.close()
    if style_xml is None:
        raise AttributeError("Нет xml файла")
    style_contains_theme = False
    # print(style_xml.tag)
    for x in style_xml:
        # print(x.tag)
        if 'style' in x.tag:
            attributes = x.attrib
            for key in attributes.keys():
                if key.endswith('styleId'):
                    style_id = attributes[key]
        for style in x:
            if style.tag.endswith('rPr'):
                for rPr in style:
                    # print(rPr.tag)
                    if rPr.tag.endswith('rFonts'):
                        rPr_attrib = rPr.attrib
                        for key in rPr_attrib.keys():
                            if key.endswith('cstheme'):
                                style_font["cstheme"] = rPr_attrib[key]
                            elif key.endswith('hAnsiTheme'):
                                style_font["hAnsiTheme"] = rPr_attrib[key]
                            elif key.endswith('eastAsiaTheme'):
                                style_font["eastAsiaTheme"] = rPr_attrib[key]
                            elif key.endswith('asciiTheme'):
                                style_font["asciiTheme"] = rPr_attrib[key]
                        if len(style_font) == 4: #возможно, правильнее проверять только тех, у кого asciiTheme и hAnsiTheme
                            styles_with_theme[style_id] = style_font
                        style_font = {}

    return styles_with_theme


def get_font_from_fonttable(file):
    fonts_fonttable = []
    with ZipFile(file, 'r') as z:
        with z.open('word/fontTable.xml') as font_table:  # todo: а я говорил
            font_table = etree.parse(font_table).getroot()
    z.close()
    if font_table is None:
        raise AttributeError("Нет xml файла")

    for x in font_table:
        font_attrib = x.attrib
        for key in font_attrib:
            if key.endswith('name'):
                fonts_fonttable.append(font_attrib[key])
    return fonts_fonttable


def correctness_fonttable(fonts_fonttable):
    """
    :param fonts_fonttable: list font name from fontTable.xml
    :return: 0 - if Times New Roman not in list,
    1 - if only Times New Roman in list, 2 - if list contains Times New Roman and other font.
    """
    if 'Times New Roman' in fonts_fonttable:
        if len(fonts_fonttable) == 1:
            return 1
        else:
            return 2
    else:
        return 3


def get_ind_content(file):
    ind_para_content = 0
    for i in range(len(file.paragraphs)):
        if file.paragraphs[i].text == "Содержание":
            ind_para_content = i
            if ind_para_content + 1 < len(file.paragraphs):
                if file.paragraphs[i + 1].style.name == "toc 1":
                    ind_para_content = i
                else:
                    ind_para_content = 0
            break
        else:
            ind_para_content = 0

    return ind_para_content


def get_field(file):
    field = False
    for section in file.sections:
        if abs(section.bottom_margin.cm - 2) < 0.001:
            field = True
        elif abs(section.top_margin.cm - 2) < 0.001:
            field = True
        elif abs(section.left_margin.cm - 3) < 0.001:
            field = True
        elif abs(section.right_margin.cm - 1) < 0.001:
            field = True
    return field


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        DOMTree = parseString(child.xml)  # TODO: тут поле для оптимизаций
        data = DOMTree.documentElement
        nodelist = data.getElementsByTagName('pic:blipFill')  # TODO: а что насчет формул? Mathtype и вордовские
        if len(nodelist) >= 1:
            for pic in nodelist:
                yield "##########################Картииинка"
        elif isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_tables_context(parent):
    para = None
    for elem in iter_block_items(parent):
        if isinstance(elem, Paragraph):
            para = elem
        elif isinstance(elem, Table):
            yield (para, elem)


def iter_pics_context(parent):
    pic = None
    for elem in iter_block_items(parent):
        if isinstance(elem, str):
            pic = elem
        elif isinstance(elem, Paragraph) and pic:
            yield (elem, pic)
            pic = None
        # else:
        #     raise Exception("Элемент не параграф и не картинка")


def tables_is_ok(parent, verifiable_paras):
    table_is_ok = False
    table_pattern = re.compile(r"Таблица \d+(\.\d+)* – .+\.")
    flag = False
    for elem in iter_tables_context(parent):
        if elem[0].text in verifiable_paras:
            flag = True
        txt = elem[0].text if elem[0] else "НЕТ ТЕКСТА"
        valid_table = re.fullmatch(table_pattern, txt)
        if valid_table:
            if flag:
                if elem[0].paragraph_format.first_line_indent == 0:
                    table_is_ok = True
                    print(txt, " \tХорошая таблица")
                elif elem[0].paragraph_format.first_line_indent is None:
                    if elem[0].style.paragraph_format.first_line_indent == 0:
                        table_is_ok = True
                        print(txt, " \tХорошая таблица")
                    else:
                        table_is_ok = False
                        print(txt, " \tПлохая таблица")
                        return table_is_ok
                else:
                    table_is_ok = False
                    print(txt, " \tПлохая таблица")
                    return table_is_ok
        else:
            if flag:
                table_is_ok = False
                print(txt, " \tПлохая таблица")
                return table_is_ok
            else:
                table_is_ok = True
                print(txt, " \tТаблица не учитывается")
    return table_is_ok


def pics_is_ok(parent, verifiable_paras):
    pic_is_ok = False
    pic_pattern = re.compile(r"Рисунок \d+(\.\d+)* – .+\.")
    flag = False
    for elem in iter_pics_context(parent):
        if elem[0].text in verifiable_paras:
            flag = True
        txt = elem[0].text if elem[0] else "НЕТ ТЕКСТА"
        valid_table = re.fullmatch(pic_pattern, txt)
        if valid_table:
            if flag:
                if elem[0].alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    print(elem[0].alignment)
                    pic_is_ok = True
                    print(txt, " \tХорошая картинка")
                elif elem[0].alignment is None:
                    if elem[0].style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        print(elem[0].style.paragraph_format.alignment)
                        pic_is_ok = True
                        print(txt, " \tХорошая картинка со стилем")
                else:
                    pic_is_ok = False
                    print(txt, " \tПлохая картинка")
                    return pic_is_ok
        else:
            if flag:
                pic_is_ok = False
                print(txt, " \tПлохая картинка")
                return pic_is_ok
            else:
                pic_is_ok = True
                print(txt, " \tКартинка не учитывается")
    return pic_is_ok


def heading_in_file(para):
    style = para.style
    heading = ['Heading 1', 'Heading 2', 'Heading 3']
    if style.name in heading:
        return True
    elif style.name is None:
        if style.base_style.name in heading:
            return True
    return False


def shorten(s):
    return f"\"{s}\"" if len(s) < 25 else f"\"{s[:25]}...\""


def text_alignment_is_ok(para):
    alignment_is_ok = False
    if para.alignment == WD_ALIGN_PARAGRAPH.JUSTYFY:
        alignment_is_ok = True
    elif para.alignment is None:
        if para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTYFY:
            alignment_is_ok = True
    else:
        alignment_is_ok = False
    return alignment_is_ok


def heading_alignment_is_ok(para):
    alignment_is_ok = False
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        alignment_is_ok = True
    elif para.alignment is None:
        if para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            alignment_is_ok = True
    else:
        alignment_is_ok = False
    return alignment_is_ok


# теперь эта функция не нужна
def get_document_font_is_ok(doc_name):
    file = docx.Document(doc_name)
    correct_font = False
    default_font_is_ok = get_default_font_is_ok(doc_name)  # Todo: нормально передавать имя файла
    ok_fonts = {"Times New Roman"}
    ind_cont = get_ind_content(file)

    for para in file.paragraphs[ind_cont:]:
        print(heading_in_file(para))
        # para-level
        para_font = None
        style = para.style
        while not para_font:
            try:
                para_font = style.font.name
                print(style.name, para_font)
                style = style.base_style
            except Exception as e:
                print(e)
                break
        if para_font not in ok_fonts:
            print(f"Paragraph-level incorrect font: {para_font} near {shorten(para.text)}")
            return False
            # para: None
            # run: написано

            # run-level
        for run in para.runs:
            font = run.font.name
            style_font = run.style.font.name
            print(font, style_font, para_font)

            if font is None:
                if style_font is None:
                    correct_font = default_font_is_ok
                else:
                    correct_font = style_font in ok_fonts

            else:
                correct_font = font in ok_fonts
                # print(f"Run-level incorrect font: {font} near {shorten(run.text)}")
                # break

            if not correct_font:
                print(f"Run-level incorrect font: {font} near {shorten(para.text)}")
                return correct_font

    return correct_font


# можно удалять верхнюю функцию

def get_font_is_ok(file_name):
    """
    :param file_name:
    :return: True if font is ok, else returns False
    """
    styles_with_theme = get_styles_with_theme(file_name)
    run_with_theme = get_theme_run(file_name)
    file = docx.Document(file_name)
    ind_content = get_ind_content(file)
    fonttable = get_font_from_fonttable(file_name)
    font_table = correctness_fonttable(fonttable)
    theme_is_ok = get_default_font_is_ok(file_name)

    ok_font = 'Times New Roman'
    font_is_ok = False
    if font_table == 0:
        print("Ошибка в таблице стилей")
        return False
    elif font_table == 1:
        font_is_ok = True
    else:
        if run_with_theme:
            if theme_is_ok:
                font_is_ok = True
            else:
                print("Ошибка в run есть тема")
                return False
        else:
            for para in file.paragraphs[ind_content:]:
                for run in para.runs:
                    if run.font.name is not None:
                        if run.font.name == ok_font:
                            font_is_ok = True
                        else:
                            print(f"Ошибка неправильный шрифт в run {shorten(para.text)}")
                            return False
                    elif run.style.font.name is not None:
                        style = run.style
                        if style.style_id in styles_with_theme:
                            if theme_is_ok:
                                font_is_ok = True
                            else:
                                print(f"Ошибка в стиле run есть тема {shorten(para.text)}")
                                return False
                        else:
                            if style.font.name == ok_font:
                                font_is_ok = True
                            else:
                                print(f"Ошибка в стиле run {shorten(para.text)}")
                                return False
                    # elif run.style.font.name is None:
                    else:
                        if run.style.style_id in styles_with_theme:
                            if theme_is_ok:
                                font_is_ok = True
                            else:
                                print(f"Ошибка в стиле run есть тема {shorten(para.text)}")
                                return False
                        else:
                            if run.style.base_style is not None:
                                font = run.style.font.name
                                style = run.style
                                while font is None:
                                    style = style.base_style
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            font_is_ok = True
                                        else:
                                            print(f"Ошибка в родительском стиле run есть тема {shorten(para.text)}")
                                            return False
                                    else:
                                        font = style.font.name
                                if font == ok_font:
                                    font_is_ok = True
                                else:
                                    print(f"Ошибка не тот шрифт в родительском стиле run {shorten(para.text)}")
                                    return False
                            else:
                                style = para.style
                                font = style.font.name
                                if font is None:
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            font_is_ok = True
                                        else:
                                            print(f"Ошибка в стиле параграфа есть тема {shorten(para.text)}")
                                            return False
                                while font is None:
                                    style = style.base_style
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            font_is_ok = True
                                        else:
                                            print(
                                                f"Ошибка в родительском стиле параграфа есть тема {shorten(para.text)}")
                                            return False
                                    else:
                                        font = style.font.name
                                if font == ok_font:
                                    font_is_ok = True
                                else:
                                    print(f"Ошибка в стиле параграфе не тот шрифт {shorten(para.text)}")
                                    return False

    return font_is_ok


file_name = "Тест.docx"
# file_name = "Курсовая работа.docx"
# file = docx.Document(file_name)
# for para in file.paragraphs:
#     for run in para.runs:
#         print()
# TODO строгий стиль ломает всё, понять, почему так
print(get_font_is_ok(file_name))

# print(get_document_font_is_ok(file))

# doc = docx.Document('Курсовая работа.docx')


# print(get_styles_with_theme('Тест.docx'))
# get_font_from_fonttable('Тест.docx')
# print(get_theme_run('Тест.docx'))

# for para in doc.paragraphs:
#     print(f'Paragraph style font {para.style.style_id}')
#     print(f'Paragraph base style font {para.style.base_style.style_id}')
#     for run in para.runs:
#         print(f'Run font {run.font.name}')
#         print(f'Run style font {run.style.style_id}')


# verifiable_paras = []
# for para in file.paragraphs[get_ind_content(file):]:
#     verifiable_paras.append(para.text)


# Если мы читаем стиль
# То мы должны проверить его XML на наличие темы. Если тема есть - то смотрим шрифты этой темы.
# Если у рана шрифт НОНЕ, то считаем, что он наследует шрифт параграфа, но нужно проверить стили

# всегда проверять наличие темы, она перекрывает стиль


# ======================================================================

# TODO Обязательно раскоментировать!
# if __name__ == "__main__":
#     import tkinter as tk
#     from tkinter import filedialog
#
#     root = tk.Tk()
#     root.withdraw()
#     file_path = filedialog.askopenfilename()
#     print(file_check_interface(file_path, mock=False))
